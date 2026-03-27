"""
PDF テキスト変換ツール - PDFからテキスト/Markdown/Wordを抽出

3エンジン構成:
  - pymupdf4llm: デジタルPDFの高速テキスト抽出（GPU不要）
  - pdfplumber:   テーブル構造を正確に抽出（表のあるPDFに最適）
  - marker-pdf:   スキャン/画像PDFの高品質OCR抽出（PyTorch使用）

出力形式:
  - txt:  プレーンテキスト
  - md:   Markdown
  - docx: Word文書（見出し・表・箇条書きを構造化）

使い方:
  CLI:  python pdf_text_tool.py input.pdf
  CLI:  python pdf_text_tool.py input.pdf -f docx
  GUI:  python pdf_text_tool.py --gui
"""

import argparse
import os
import re
import sys
import threading
import time
from pathlib import Path


# ---------------------------------------------------------------------------
# バリデーション
# ---------------------------------------------------------------------------

def validate_pdf_path(input_path):
    """入力ファイルのバリデーション。問題があればエラーメッセージを返す。"""
    input_path = Path(input_path)
    if not input_path.exists():
        return f"ファイルが見つかりません: {input_path}"
    if input_path.suffix.lower() != ".pdf":
        return f"PDFファイルではありません: {input_path.name} (拡張子: {input_path.suffix})"
    return None


def is_pdf_encrypted(input_path):
    """PDFがパスワード保護されているか確認"""
    try:
        import pymupdf
        doc = pymupdf.open(str(input_path))
        encrypted = doc.is_encrypted
        doc.close()
        return encrypted
    except Exception:
        return False


# ---------------------------------------------------------------------------
# PDF種別判定
# ---------------------------------------------------------------------------

def detect_pdf_type(input_path):
    """
    PDFがデジタル（テキスト埋め込み）かスキャン（画像ベース）かを判定。
    """
    try:
        import pymupdf
    except ImportError:
        return "unknown"

    try:
        doc = pymupdf.open(str(input_path))
    except Exception:
        return "unknown"

    total_pages = len(doc)
    if total_pages == 0:
        doc.close()
        return "unknown"

    pages_with_text = 0
    total_text_len = 0
    sample_pages = min(total_pages, 10)

    for i in range(sample_pages):
        page = doc[i]
        text = page.get_text().strip()
        text_len = len(text)
        total_text_len += text_len
        if text_len > 10:
            pages_with_text += 1

    doc.close()

    if pages_with_text > 0 or total_text_len > 20:
        return "digital"
    return "scanned"


# ---------------------------------------------------------------------------
# 構造化抽出（pdfplumber ベース）
# ---------------------------------------------------------------------------

def extract_structured(input_path, progress_callback=None):
    """
    pdfplumber を使って PDF の構造（テキスト・テーブル）を抽出する。

    Returns:
        list[dict]: ページごとの構造化データ
        [
            {
                "page": 1,
                "text": "ページ全体のテキスト",
                "tables": [
                    [["col1", "col2"], ["val1", "val2"], ...]
                ]
            },
            ...
        ]
    """
    def log(msg):
        if progress_callback:
            progress_callback(msg)

    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError(
            "pdfplumber がインストールされていません。\n"
            "  → pip install pdfplumber を実行してください。"
        )

    log("pdfplumber で構造化テキスト抽出中...")

    pages_data = []
    with pdfplumber.open(str(input_path)) as pdf:
        total = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            if total > 5 and (i + 1) % 5 == 0:
                log(f"  ページ {i+1}/{total} 処理中...")

            text = page.extract_text() or ""
            tables = page.extract_tables() or []

            # テーブル内のNoneを空文字に変換
            clean_tables = []
            for table in tables:
                clean_table = []
                for row in table:
                    clean_table.append([
                        (cell or "").strip() for cell in row
                    ])
                clean_tables.append(clean_table)

            pages_data.append({
                "page": i + 1,
                "text": text.strip(),
                "tables": clean_tables,
            })

    log(f"  {len(pages_data)} ページ抽出完了")
    return pages_data


# ---------------------------------------------------------------------------
# テキスト抽出エンジン
# ---------------------------------------------------------------------------

def extract_text_pymupdf(input_path, output_format="txt", progress_callback=None):
    """pymupdf4llm を使用してデジタルPDFからテキスト抽出。"""
    def log(msg):
        if progress_callback:
            progress_callback(msg)

    try:
        import pymupdf4llm
    except ImportError:
        raise RuntimeError(
            "pymupdf4llm がインストールされていません。\n"
            "  → pip install pymupdf4llm を実行してください。"
        )

    if output_format == "txt":
        log("pymupdf4llm.to_text() でテキスト抽出中...")
        text = pymupdf4llm.to_text(str(input_path), show_progress=False)
        log("テキスト変換完了")
        return text.strip()

    log("pymupdf4llm.to_markdown() でMarkdown抽出中...")
    md_text = pymupdf4llm.to_markdown(
        str(input_path), ignore_images=True, show_progress=False
    )
    log("Markdown変換完了")
    return md_text.strip()


def extract_text_marker(input_path, output_format="txt", progress_callback=None):
    """marker-pdf を使用してスキャン/画像PDFから高品質OCRテキスト抽出。"""
    def log(msg):
        if progress_callback:
            progress_callback(msg)

    try:
        from marker.converters.pdf import PdfConverter
        from marker.models import create_model_dict
        from marker.output import text_from_rendered
    except ImportError:
        raise RuntimeError(
            "marker-pdf がインストールされていません。\n"
            "  → pip install marker-pdf を実行してください。\n"
            "  （PyTorchも必要です）"
        )

    log("marker-pdf でOCRテキスト抽出中（時間がかかる場合があります）...")

    converter = PdfConverter(artifact_dict=create_model_dict())
    rendered = converter(str(input_path))
    text, _, images = text_from_rendered(rendered)

    if output_format == "txt":
        text = _markdown_to_plain_text(text)
        log("テキスト変換完了")
        return text

    log("Markdown変換完了")
    return text


def _markdown_to_plain_text(md_text):
    """Markdownテキストからプレーンテキストに変換（marker-pdf出力用）"""
    text = md_text
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'```[^\n]*\n(.*?)```', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'^\|[-:| ]+\|\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(
        r'^\|(.+)\|\s*$',
        lambda m: '\t'.join(cell.strip() for cell in m.group(1).split('|')),
        text, flags=re.MULTILINE
    )
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'^[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^>\s?', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ---------------------------------------------------------------------------
# Word (.docx) 出力
# ---------------------------------------------------------------------------

def save_as_docx(pages_data, output_path, source_filename="", progress_callback=None):
    """
    構造化データからWord文書を生成。

    Args:
        pages_data: extract_structured() の戻り値
        output_path: 出力 .docx パス
        source_filename: 元PDFファイル名（ヘッダー用）
        progress_callback: fn(message: str)
    """
    def log(msg):
        if progress_callback:
            progress_callback(msg)

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise RuntimeError(
            "python-docx がインストールされていません。\n"
            "  → pip install python-docx を実行してください。"
        )

    log("Word文書を生成中...")

    doc = Document()

    # デフォルトスタイル設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(10.5)

    # ソースファイル名をヘッダーに
    if source_filename:
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"変換元: {source_filename}"
        header_para.style.font.size = Pt(8)
        header_para.style.font.color.rgb = RGBColor(128, 128, 128)

    for page_data in pages_data:
        text = page_data["text"]
        tables = page_data["tables"]

        if not text and not tables:
            continue

        # テーブルのテキストを集約（本文テキストからテーブル部分を識別するため）
        table_cell_texts = set()
        for table in tables:
            for row in table:
                for cell in row:
                    if cell and len(cell) > 1:
                        table_cell_texts.add(cell.strip())

        # テキストを行ごとに処理
        lines = text.split('\n') if text else []
        i = 0
        pending_table_idx = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # テーブル行の検出: テーブルセルと一致する行が連続したらテーブルを挿入
            if (pending_table_idx < len(tables) and
                    _line_matches_table(line, tables[pending_table_idx], table_cell_texts)):
                # テーブル挿入
                table_data = tables[pending_table_idx]
                _add_table_to_doc(doc, table_data)
                pending_table_idx += 1

                # テーブルに対応するテキスト行をスキップ
                while i < len(lines):
                    l = lines[i].strip()
                    if l and any(cell in l for cell in table_cell_texts if cell):
                        i += 1
                    elif not l:
                        i += 1
                    else:
                        break
                continue

            # 見出しパターンの検出
            heading_match = re.match(
                r'^(\d+)\.\s+(.+)$', line
            )
            if heading_match:
                level = 1 if int(heading_match.group(1)) < 10 else 2
                doc.add_heading(line, level=level)
                i += 1
                continue

            # 箇条書きパターンの検出
            bullet_match = re.match(r'^[●○■□▪▸▹・\-\*]\s*(.+)$', line)
            if bullet_match:
                para = doc.add_paragraph(bullet_match.group(1), style='List Bullet')
                i += 1
                continue

            # 通常の段落
            doc.add_paragraph(line)
            i += 1

        # 残りのテーブルを挿入
        while pending_table_idx < len(tables):
            _add_table_to_doc(doc, tables[pending_table_idx])
            pending_table_idx += 1

        # ページ区切り（最後のページ以外）
        if page_data["page"] < len(pages_data):
            doc.add_page_break()

    doc.save(str(output_path))
    log(f"Word文書保存完了: {Path(output_path).name}")


def _line_matches_table(line, table, table_cell_texts):
    """行がテーブルのヘッダー行に含まれるかチェック"""
    if not table or not table[0]:
        return False
    header_cells = [c for c in table[0] if c]
    if not header_cells:
        return False
    match_count = sum(1 for cell in header_cells if cell in line)
    return match_count >= max(1, len(header_cells) // 2)


def _add_table_to_doc(doc, table_data):
    """Word文書にテーブルを追加"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    if not table_data or not table_data[0]:
        return

    num_cols = max(len(row) for row in table_data)
    num_rows = len(table_data)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    for r, row_data in enumerate(table_data):
        row = table.rows[r]
        for c, cell_text in enumerate(row_data):
            if c < num_cols:
                cell = row.cells[c]
                cell.text = cell_text or ""
                # セルのフォント設定
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Yu Gothic'

    # ヘッダー行を太字に
    if table_data:
        for c in range(num_cols):
            cell = table.rows[0].cells[c]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    doc.add_paragraph()  # テーブル後にスペース


# ---------------------------------------------------------------------------
# 統合抽出関数
# ---------------------------------------------------------------------------

def extract_text(input_path, engine="auto", output_format="txt",
                 progress_callback=None):
    """PDFからテキストを抽出する統合関数。"""
    input_path = Path(input_path)

    err = validate_pdf_path(input_path)
    if err:
        raise ValueError(err)

    if is_pdf_encrypted(input_path):
        raise ValueError(f"パスワード保護されたPDFです: {input_path.name}\n"
                         "  パスワードを解除してから再度お試しください。")

    def log(msg):
        if progress_callback:
            progress_callback(msg)

    log(f"処理開始: {input_path.name}")

    # エンジン選択
    if engine == "auto":
        log("PDF種別を判定中...")
        pdf_type = detect_pdf_type(input_path)
        log(f"  判定結果: {pdf_type}")

        if pdf_type == "scanned":
            try:
                import marker  # noqa: F401
                engine = "marker"
                log("  → marker-pdf (高精度OCR) を使用します")
            except ImportError:
                log("  → marker-pdf 未インストール。pymupdf4llm で試行します")
                engine = "pymupdf"
        else:
            engine = "pymupdf"
            log("  → pymupdf4llm (高速) を使用します")

    start_time = time.time()

    if engine == "marker":
        result = extract_text_marker(input_path, output_format, progress_callback)
    else:
        result = extract_text_pymupdf(input_path, output_format, progress_callback)

    elapsed = time.time() - start_time
    log(f"  処理時間: {elapsed:.1f}秒")

    if not result or not result.strip():
        log("  警告: テキストが抽出できませんでした（空の結果）")

    return result


def extract_and_save(input_path, output_path=None, engine="auto",
                     output_format="txt", progress_callback=None):
    """PDFからテキストを抽出してファイルに保存。"""
    input_path = Path(input_path)

    # バリデーション
    err = validate_pdf_path(input_path)
    if err:
        raise ValueError(err)

    if is_pdf_encrypted(input_path):
        raise ValueError(f"パスワード保護されたPDFです: {input_path.name}\n"
                         "  パスワードを解除してから再度お試しください。")

    # 出力パス決定
    ext_map = {"txt": ".txt", "md": ".md", "docx": ".docx"}
    ext = ext_map.get(output_format, ".txt")

    if output_path is None:
        output_path = input_path.with_suffix(ext)
    output_path = Path(output_path)

    # 入力と出力が同一ファイルになる場合の防止
    try:
        if input_path.resolve() == output_path.resolve():
            stem = input_path.stem + "_extracted"
            output_path = input_path.parent / (stem + ext)
    except (OSError, ValueError):
        pass

    def log(msg):
        if progress_callback:
            progress_callback(msg)

    log(f"処理開始: {input_path.name}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    start_time = time.time()

    if output_format == "docx":
        # Word出力: pdfplumber で構造化抽出 → docx 生成
        pages_data = extract_structured(input_path, progress_callback)
        save_as_docx(
            pages_data, output_path,
            source_filename=input_path.name,
            progress_callback=progress_callback
        )
    else:
        # テキスト/Markdown出力
        text = extract_text(
            input_path, engine=engine, output_format=output_format,
            progress_callback=progress_callback
        )
        output_path.write_text(text, encoding="utf-8")

        char_count = len(text)
        line_count = text.count('\n') + 1 if text else 0
        log(f"保存完了: {output_path.name} ({char_count:,}文字, {line_count:,}行)")

    elapsed = time.time() - start_time
    log(f"  総処理時間: {elapsed:.1f}秒")

    return str(output_path)


# ---------------------------------------------------------------------------
# 依存関係チェック
# ---------------------------------------------------------------------------

def check_dependencies():
    """利用可能なエンジンと出力形式を確認"""
    results = {}

    try:
        import pymupdf4llm  # noqa: F401
        results["pymupdf4llm"] = True
    except ImportError:
        results["pymupdf4llm"] = False

    try:
        import pdfplumber  # noqa: F401
        results["pdfplumber"] = True
    except ImportError:
        results["pdfplumber"] = False

    try:
        import marker  # noqa: F401
        results["marker"] = True
    except ImportError:
        results["marker"] = False

    try:
        import docx  # noqa: F401
        results["python-docx"] = True
    except ImportError:
        results["python-docx"] = False

    return results


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def run_cli():
    parser = argparse.ArgumentParser(
        description="PDF テキスト変換ツール - PDFからテキスト/Markdown/Wordを抽出"
    )
    parser.add_argument(
        "input", nargs="?",
        help="入力PDFファイルパス"
    )
    parser.add_argument(
        "-o", "--output",
        help="出力ファイルパス (省略時: <入力名>.txt/.md/.docx)"
    )
    parser.add_argument(
        "-e", "--engine", default="auto",
        choices=["auto", "pymupdf", "marker"],
        help="抽出エンジン (デフォルト: auto)"
    )
    parser.add_argument(
        "-f", "--format", default="txt",
        choices=["txt", "md", "docx"],
        help="出力形式 (デフォルト: txt)"
    )
    parser.add_argument(
        "--gui", action="store_true",
        help="GUIモードで起動"
    )
    parser.add_argument(
        "--check", action="store_true",
        help="依存関係を確認して終了"
    )

    args = parser.parse_args()

    if args.check:
        deps = check_dependencies()
        print("=== 依存関係チェック ===")
        for name, available in deps.items():
            status = "OK" if available else "未インストール"
            print(f"  {name}: {status}")
        can_extract = deps.get("pymupdf4llm") or deps.get("pdfplumber")
        if not can_extract:
            print("\nエラー: 少なくとも1つの抽出エンジンが必要です。")
            print("  → pip install pymupdf4llm pdfplumber")
            sys.exit(1)
        else:
            print("\n利用可能です。")
            sys.exit(0)

    if args.gui or args.input is None:
        run_gui()
        return

    err = validate_pdf_path(args.input)
    if err:
        print(f"エラー: {err}", file=sys.stderr)
        sys.exit(1)

    deps = check_dependencies()
    if args.format == "docx" and not deps.get("python-docx"):
        print("エラー: Word出力にはpython-docxが必要です。", file=sys.stderr)
        print("  → pip install python-docx", file=sys.stderr)
        sys.exit(1)

    if args.format == "docx" and not deps.get("pdfplumber"):
        print("エラー: Word出力にはpdfplumberが必要です。", file=sys.stderr)
        print("  → pip install pdfplumber", file=sys.stderr)
        sys.exit(1)

    try:
        output = extract_and_save(
            args.input,
            args.output,
            engine=args.engine,
            output_format=args.format,
            progress_callback=lambda msg: print(f"  {msg}")
        )
        print(f"\n出力: {output}")
    except Exception as e:
        print(f"\nエラー: {e}", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------

def run_gui():
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox

    has_windnd = False
    try:
        import windnd
        has_windnd = True
    except ImportError:
        pass

    class App:
        def __init__(self, root):
            self.root = root
            self.root.title("PDF テキスト変換ツール")
            self.root.geometry("720x680")
            self.root.resizable(True, True)
            self.root.configure(bg="#f0f0f0")

            self.files = []
            self.processing = False

            self._build_ui()
            self._check_deps()

            if has_windnd:
                windnd.hook_dropfiles(self.root, func=self._on_drop)

        def _build_ui(self):
            title_frame = tk.Frame(self.root, bg="#1a5276", pady=10)
            title_frame.pack(fill=tk.X)
            tk.Label(
                title_frame, text="PDF テキスト変換ツール",
                font=("Segoe UI", 16, "bold"), fg="white", bg="#1a5276"
            ).pack()
            tk.Label(
                title_frame,
                text="PDFからテキスト/Markdown/Wordを抽出",
                font=("Segoe UI", 9), fg="#aed6f1", bg="#1a5276"
            ).pack()

            # ドロップエリア
            drop_frame = tk.LabelFrame(
                self.root, text=" PDFファイル ",
                font=("Segoe UI", 10), padx=10, pady=5, bg="#f0f0f0"
            )
            drop_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(10, 5))

            self.drop_label = tk.Label(
                drop_frame,
                text=("ここにPDFをドラッグ&ドロップ\nまたは下のボタンでファイルを選択"
                      if has_windnd else
                      "下のボタンでファイルを選択してください"),
                font=("Segoe UI", 11),
                bg="#eaf2f8", fg="#7f8c8d",
                relief=tk.GROOVE, bd=2,
                width=50, height=3
            )
            self.drop_label.pack(fill=tk.BOTH, expand=True, pady=(5, 5))

            self.file_listbox = tk.Listbox(
                drop_frame, height=4,
                font=("Segoe UI", 9), selectmode=tk.EXTENDED
            )
            self.file_listbox.pack(fill=tk.BOTH, expand=True, pady=(0, 5))

            btn_frame = tk.Frame(drop_frame, bg="#f0f0f0")
            btn_frame.pack(fill=tk.X)
            ttk.Button(btn_frame, text="ファイルを追加...",
                       command=self._add_files).pack(side=tk.LEFT, padx=(0, 5))
            ttk.Button(btn_frame, text="選択を削除",
                       command=self._remove_selected).pack(side=tk.LEFT, padx=(0, 5))
            ttk.Button(btn_frame, text="すべてクリア",
                       command=self._clear_files).pack(side=tk.LEFT)

            # 設定
            settings_frame = tk.LabelFrame(
                self.root, text=" 設定 ",
                font=("Segoe UI", 10), padx=10, pady=5, bg="#f0f0f0"
            )
            settings_frame.pack(fill=tk.X, padx=15, pady=5)

            tk.Label(settings_frame, text="エンジン:",
                     font=("Segoe UI", 10), bg="#f0f0f0"
                     ).grid(row=0, column=0, sticky=tk.W, pady=2)

            self.engine_var = tk.StringVar(value="auto (自動判定)")
            ttk.Combobox(
                settings_frame, textvariable=self.engine_var,
                values=[
                    "auto (自動判定)",
                    "pymupdf (高速・デジタルPDF向け)",
                    "marker (高精度OCR・スキャンPDF向け)"
                ],
                width=35, state="readonly"
            ).grid(row=0, column=1, sticky=tk.W, padx=10, pady=2)

            tk.Label(settings_frame, text="出力形式:",
                     font=("Segoe UI", 10), bg="#f0f0f0"
                     ).grid(row=1, column=0, sticky=tk.W, pady=2)

            self.format_var = tk.StringVar(value="docx (Word文書)")
            ttk.Combobox(
                settings_frame, textvariable=self.format_var,
                values=[
                    "docx (Word文書)",
                    "txt (プレーンテキスト)",
                    "md (Markdown)",
                ],
                width=35, state="readonly"
            ).grid(row=1, column=1, sticky=tk.W, padx=10, pady=2)

            tk.Label(settings_frame, text="出力先:",
                     font=("Segoe UI", 10), bg="#f0f0f0"
                     ).grid(row=2, column=0, sticky=tk.W, pady=2)

            output_row = tk.Frame(settings_frame, bg="#f0f0f0")
            output_row.grid(row=2, column=1, sticky=tk.W, padx=10, pady=2)

            self.output_var = tk.StringVar(value="same")
            ttk.Radiobutton(
                output_row, text="入力と同じフォルダ",
                variable=self.output_var, value="same",
                command=self._on_output_mode_change
            ).pack(anchor=tk.W)

            custom_row = tk.Frame(output_row, bg="#f0f0f0")
            custom_row.pack(anchor=tk.W, fill=tk.X)
            ttk.Radiobutton(
                custom_row, text="フォルダを指定:",
                variable=self.output_var, value="custom",
                command=self._on_output_mode_change
            ).pack(side=tk.LEFT)

            self.output_dir_var = tk.StringVar(value="")
            self.output_dir_entry = ttk.Entry(
                custom_row, textvariable=self.output_dir_var, width=25,
                state=tk.DISABLED
            )
            self.output_dir_entry.pack(side=tk.LEFT, padx=(5, 3))

            self.output_dir_btn = ttk.Button(
                custom_row, text="参照...",
                command=self._browse_output_dir, state=tk.DISABLED
            )
            self.output_dir_btn.pack(side=tk.LEFT)

            # 実行ボタン
            exec_frame = tk.Frame(self.root, bg="#f0f0f0", pady=5)
            exec_frame.pack(fill=tk.X, padx=15)

            self.run_btn = ttk.Button(
                exec_frame, text="変換実行",
                command=self._start_extraction
            )
            self.run_btn.pack(fill=tk.X, ipady=5)

            self.progress = ttk.Progressbar(self.root, mode="indeterminate")
            self.progress.pack(fill=tk.X, padx=15, pady=(5, 0))

            # ログ
            log_frame = tk.LabelFrame(
                self.root, text=" ログ ",
                font=("Segoe UI", 10), padx=5, pady=5, bg="#f0f0f0"
            )
            log_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=(5, 10))

            self.log_text = tk.Text(
                log_frame, height=6,
                font=("Consolas", 9), state=tk.DISABLED, wrap=tk.WORD
            )
            scrollbar = ttk.Scrollbar(log_frame, command=self.log_text.yview)
            self.log_text.configure(yscrollcommand=scrollbar.set)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            self.log_text.pack(fill=tk.BOTH, expand=True)

            self.status_var = tk.StringVar(value="準備完了")
            tk.Label(
                self.root, textvariable=self.status_var,
                font=("Segoe UI", 9), fg="#7f8c8d", bg="#f0f0f0",
                anchor=tk.W
            ).pack(fill=tk.X, padx=15, pady=(0, 5))

        def _check_deps(self):
            deps = check_dependencies()
            available = [n for n, ok in deps.items() if ok]
            missing = [n for n, ok in deps.items() if not ok]
            if available:
                self._log(f"利用可能: {', '.join(available)}")
            if missing:
                self._log(f"未インストール: {', '.join(missing)}")
            if not any(deps.get(e) for e in ["pymupdf4llm", "pdfplumber"]):
                self._log("エラー: エンジンが1つもインストールされていません！")
                self.status_var.set("エンジン未インストール")

        def _log(self, msg):
            self.log_text.configure(state=tk.NORMAL)
            self.log_text.insert(tk.END, msg + "\n")
            self.log_text.see(tk.END)
            self.log_text.configure(state=tk.DISABLED)

        def _on_output_mode_change(self):
            if self.output_var.get() == "custom":
                self.output_dir_entry.configure(state=tk.NORMAL)
                self.output_dir_btn.configure(state=tk.NORMAL)
                if not self.output_dir_var.get():
                    self._browse_output_dir()
            else:
                self.output_dir_entry.configure(state=tk.DISABLED)
                self.output_dir_btn.configure(state=tk.DISABLED)

        def _browse_output_dir(self):
            d = filedialog.askdirectory(title="出力先フォルダを選択")
            if d:
                self.output_dir_var.set(d)

        def _on_drop(self, files):
            for f in files:
                if isinstance(f, bytes):
                    for enc in ("utf-8", "cp932", "shift_jis"):
                        try:
                            f = f.decode(enc)
                            break
                        except (UnicodeDecodeError, LookupError):
                            continue
                    else:
                        f = f.decode("utf-8", errors="replace")
                f = str(f).strip().strip('"')
                if f.lower().endswith(".pdf") and f not in self.files:
                    self.files.append(f)
                    self.file_listbox.insert(tk.END, os.path.basename(f))
                    self._log(f"追加: {os.path.basename(f)}")
            self._update_drop_label()

        def _add_files(self):
            paths = filedialog.askopenfilenames(
                title="PDFファイルを選択",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
            )
            for p in paths:
                if p not in self.files:
                    self.files.append(p)
                    self.file_listbox.insert(tk.END, os.path.basename(p))
            self._update_drop_label()

        def _remove_selected(self):
            for i in reversed(list(self.file_listbox.curselection())):
                self.file_listbox.delete(i)
                del self.files[i]
            self._update_drop_label()

        def _clear_files(self):
            self.files.clear()
            self.file_listbox.delete(0, tk.END)
            self._update_drop_label()

        def _update_drop_label(self):
            count = len(self.files)
            if count > 0:
                self.drop_label.configure(
                    text=f"{count} 個のPDFファイルが選択されています",
                    fg="#27ae60"
                )
            else:
                self.drop_label.configure(
                    text=("ここにPDFをドラッグ&ドロップ\n"
                          "または下のボタンでファイルを選択"
                          if has_windnd else
                          "下のボタンでファイルを選択してください"),
                    fg="#7f8c8d"
                )

        def _get_engine(self):
            val = self.engine_var.get()
            if val.startswith("pymupdf"):
                return "pymupdf"
            elif val.startswith("marker"):
                return "marker"
            return "auto"

        def _get_format(self):
            val = self.format_var.get()
            if val.startswith("docx"):
                return "docx"
            elif val.startswith("md"):
                return "md"
            return "txt"

        def _start_extraction(self):
            if self.processing:
                return
            if not self.files:
                messagebox.showwarning("警告", "PDFファイルを選択してください。")
                return

            deps = check_dependencies()
            fmt = self._get_format()

            if fmt == "docx":
                if not deps.get("pdfplumber"):
                    messagebox.showerror("エラー",
                                         "Word出力にはpdfplumberが必要です。\n"
                                         "pip install pdfplumber を実行してください。")
                    return
                if not deps.get("python-docx"):
                    messagebox.showerror("エラー",
                                         "Word出力にはpython-docxが必要です。\n"
                                         "pip install python-docx を実行してください。")
                    return
            else:
                if not any(deps.get(e) for e in ["pymupdf4llm", "pdfplumber"]):
                    messagebox.showerror("エラー",
                                         "エンジンがインストールされていません。\n"
                                         "pip install pymupdf4llm を実行してください。")
                    return

            if self.output_var.get() == "custom":
                out_dir = self.output_dir_var.get()
                if not out_dir or not os.path.isdir(out_dir):
                    messagebox.showwarning("警告", "出力先フォルダを選択してください。")
                    return

            self.processing = True
            self.run_btn.configure(state=tk.DISABLED)
            self.progress.start(10)

            thread = threading.Thread(target=self._extraction_worker, daemon=True)
            thread.start()

        def _extraction_worker(self):
            engine = self._get_engine()
            out_format = self._get_format()
            output_mode = self.output_var.get()
            output_dir = (self.output_dir_var.get()
                          if output_mode == "custom" else None)
            total = len(self.files)
            success = 0
            failed = 0
            overall_start = time.time()

            ext_map = {"txt": ".txt", "md": ".md", "docx": ".docx"}
            ext = ext_map.get(out_format, ".txt")

            for i, filepath in enumerate(self.files):
                self.root.after(0, self.status_var.set,
                                f"処理中: {i+1}/{total}")

                if output_dir:
                    fname = Path(filepath).stem + ext
                    out_path = str(Path(output_dir) / fname)
                else:
                    out_path = str(Path(filepath).with_suffix(ext))

                try:
                    extract_and_save(
                        filepath,
                        output_path=out_path,
                        engine=engine,
                        output_format=out_format,
                        progress_callback=lambda msg: self.root.after(
                            0, self._log, msg
                        )
                    )
                    success += 1
                except Exception as e:
                    self.root.after(0, self._log, f"エラー: {e}")
                    failed += 1

            overall_elapsed = time.time() - overall_start

            def finish():
                self.progress.stop()
                self.run_btn.configure(state=tk.NORMAL)
                self.processing = False
                self.status_var.set(
                    f"完了: 成功 {success} / 失敗 {failed} / "
                    f"合計 {total} ({overall_elapsed:.1f}秒)"
                )
                self._log(f"全体処理時間: {overall_elapsed:.1f}秒")
                messagebox.showinfo(
                    "完了",
                    f"変換が完了しました。\n"
                    f"成功: {success}\n失敗: {failed}\n合計: {total}\n"
                    f"処理時間: {overall_elapsed:.1f}秒"
                )

            self.root.after(0, finish)

    root = tk.Tk()
    app = App(root)  # noqa: F841
    root.mainloop()


# ---------------------------------------------------------------------------
# エントリポイント
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    run_cli()
